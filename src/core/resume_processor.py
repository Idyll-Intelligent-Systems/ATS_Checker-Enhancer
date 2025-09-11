"""
ZeX-ATS-AI Resume Processor
Handles resume parsing from multiple formats and text extraction.
"""

import io
import re
from typing import Dict, List, Optional, BinaryIO, Union
from pathlib import Path
import asyncio
from dataclasses import dataclass

import PyPDF2
import pdfplumber
from docx import Document
from minio import Minio

from src.core.config import settings
from src.utils.text_cleaner import TextCleaner
from src.utils.file_validator import FileValidator


@dataclass
class DocumentMetadata:
    """Document metadata information."""
    filename: str
    file_size: int
    file_type: str
    page_count: int
    word_count: int
    character_count: int
    extraction_method: str
    processing_time: float
    
    def to_dict(self) -> Dict:
        return {
            'filename': self.filename,
            'file_size': self.file_size,
            'file_type': self.file_type,
            'page_count': self.page_count,
            'word_count': self.word_count,
            'character_count': self.character_count,
            'extraction_method': self.extraction_method,
            'processing_time': self.processing_time
        }


@dataclass 
class ProcessedResume:
    """Processed resume with extracted content and metadata."""
    text_content: str
    cleaned_text: str
    sections: Dict[str, str]
    metadata: DocumentMetadata
    contact_info: Dict[str, Optional[str]]
    
    def to_dict(self) -> Dict:
        return {
            'text_content': self.text_content,
            'cleaned_text': self.cleaned_text,
            'sections': self.sections,
            'metadata': self.metadata.to_dict(),
            'contact_info': self.contact_info
        }


class ResumeProcessor:
    """Enterprise resume processing with multiple format support."""
    
    def __init__(self):
        """Initialize the resume processor."""
        self.text_cleaner = TextCleaner()
        self.file_validator = FileValidator()
        
        # Initialize MinIO client for file storage
        self.minio_client = Minio(
            settings.minio_endpoint,
            access_key=settings.minio_access_key,
            secret_key=settings.minio_secret_key,
            secure=False  # Set to True for HTTPS
        )
        
        # Ensure bucket exists
        self._ensure_bucket_exists()
    
    def _ensure_bucket_exists(self):
        """Ensure MinIO bucket exists for document storage."""
        try:
            if not self.minio_client.bucket_exists(settings.minio_bucket_name):
                self.minio_client.make_bucket(settings.minio_bucket_name)
        except Exception as e:
            print(f"Warning: Could not create MinIO bucket: {e}")
    
    async def process_resume_file(
        self, 
        file_content: BinaryIO, 
        filename: str,
        store_file: bool = True
    ) -> ProcessedResume:
        """
        Process resume file and extract structured information.
        
        Args:
            file_content: Binary file content
            filename: Original filename
            store_file: Whether to store file in object storage
            
        Returns:
            ProcessedResume object with extracted content and metadata
        """
        import time
        start_time = time.time()
        
        # Validate file
        validation_result = await self.file_validator.validate_file(file_content, filename)
        if not validation_result.is_valid:
            raise ValueError(f"File validation failed: {validation_result.error_message}")
        
        # Store file if requested
        if store_file:
            await self._store_file(file_content, filename)
        
        # Extract text based on file type
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            text_content, extraction_method, page_count = await self._extract_from_pdf(file_content)
        elif file_extension in ['.doc', '.docx']:
            text_content, extraction_method, page_count = await self._extract_from_word(file_content)
        elif file_extension == '.txt':
            text_content, extraction_method, page_count = await self._extract_from_text(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Clean and structure the text
        cleaned_text = self.text_cleaner.clean_text(text_content)
        
        # Extract sections
        sections = await self._extract_sections(cleaned_text)
        
        # Extract contact information
        contact_info = await self._extract_contact_info(cleaned_text)
        
        # Create metadata
        processing_time = time.time() - start_time
        metadata = DocumentMetadata(
            filename=filename,
            file_size=len(file_content.read()),
            file_type=file_extension,
            page_count=page_count,
            word_count=len(cleaned_text.split()),
            character_count=len(cleaned_text),
            extraction_method=extraction_method,
            processing_time=processing_time
        )
        
        # Reset file pointer
        file_content.seek(0)
        
        return ProcessedResume(
            text_content=text_content,
            cleaned_text=cleaned_text,
            sections=sections,
            metadata=metadata,
            contact_info=contact_info
        )
    
    async def process_resume_text(self, text: str, filename: str = "text_input.txt") -> ProcessedResume:
        """
        Process resume from plain text input.
        
        Args:
            text: Resume text content
            filename: Virtual filename for metadata
            
        Returns:
            ProcessedResume object
        """
        import time
        start_time = time.time()
        
        # Clean text
        cleaned_text = self.text_cleaner.clean_text(text)
        
        # Extract sections and contact info
        sections = await self._extract_sections(cleaned_text)
        contact_info = await self._extract_contact_info(cleaned_text)
        
        # Create metadata
        processing_time = time.time() - start_time
        metadata = DocumentMetadata(
            filename=filename,
            file_size=len(text.encode('utf-8')),
            file_type='.txt',
            page_count=1,
            word_count=len(cleaned_text.split()),
            character_count=len(cleaned_text),
            extraction_method='direct_text',
            processing_time=processing_time
        )
        
        return ProcessedResume(
            text_content=text,
            cleaned_text=cleaned_text,
            sections=sections,
            metadata=metadata,
            contact_info=contact_info
        )
    
    async def _extract_from_pdf(self, file_content: BinaryIO) -> tuple[str, str, int]:
        """Extract text from PDF file using multiple methods."""
        file_content.seek(0)
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_content) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    return text, "pdfplumber", len(pdf.pages)
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyPDF2
        file_content.seek(0)
        try:
            reader = PyPDF2.PdfReader(file_content)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                return text, "PyPDF2", len(reader.pages)
        except Exception as e:
            print(f"PyPDF2 extraction failed: {e}")
        
        raise ValueError("Could not extract text from PDF file")
    
    async def _extract_from_word(self, file_content: BinaryIO) -> tuple[str, str, int]:
        """Extract text from Word document."""
        file_content.seek(0)
        
        try:
            doc = Document(file_content)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Estimate page count (roughly 500 words per page)
            word_count = len(text.split())
            page_count = max(1, word_count // 500)
            
            return text, "python-docx", page_count
        
        except Exception as e:
            raise ValueError(f"Could not extract text from Word document: {e}")
    
    async def _extract_from_text(self, file_content: BinaryIO) -> tuple[str, str, int]:
        """Extract text from plain text file."""
        file_content.seek(0)
        
        try:
            text = file_content.read().decode('utf-8')
            return text, "direct_read", 1
        except UnicodeDecodeError:
            # Try different encodings
            file_content.seek(0)
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_content.read().decode(encoding)
                    return text, f"direct_read_{encoding}", 1
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
    
    async def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract resume sections using pattern matching."""
        sections = {}
        
        # Define section patterns
        section_patterns = {
            'summary': r'(summary|profile|objective)[\s]*:?(.*?)(?=\n\s*(?:experience|education|skills|projects|certifications|\Z))',
            'experience': r'(experience|work\s+history|employment)[\s]*:?(.*?)(?=\n\s*(?:education|skills|projects|certifications|\Z))',
            'education': r'(education|academic|qualifications)[\s]*:?(.*?)(?=\n\s*(?:experience|skills|projects|certifications|\Z))',
            'skills': r'(skills|competencies|technologies)[\s]*:?(.*?)(?=\n\s*(?:experience|education|projects|certifications|\Z))',
            'projects': r'(projects|portfolio)[\s]*:?(.*?)(?=\n\s*(?:experience|education|skills|certifications|\Z))',
            'certifications': r'(certifications|certificates|licenses)[\s]*:?(.*?)(?=\n\s*(?:experience|education|skills|projects|\Z))'
        }
        
        text_lower = text.lower()
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text_lower, re.IGNORECASE | re.DOTALL)
            if match:
                # Get the actual text with original case
                start_pos = match.start(2)
                end_pos = match.end(2)
                section_text = text[start_pos:end_pos].strip()
                if section_text:
                    sections[section_name] = section_text
        
        return sections
    
    async def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume text."""
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'website': None,
            'location': None
        }
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone extraction
        phone_patterns = [
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(\d{3}\)\s?\d{3}-?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info['phone'] = phone_match.group().strip()
                break
        
        # LinkedIn extraction
        linkedin_patterns = [
            r'linkedin\.com/in/[A-Za-z0-9-]+',
            r'www\.linkedin\.com/in/[A-Za-z0-9-]+',
            r'https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9-]+'
        ]
        
        for pattern in linkedin_patterns:
            linkedin_match = re.search(pattern, text, re.IGNORECASE)
            if linkedin_match:
                contact_info['linkedin'] = linkedin_match.group()
                break
        
        # GitHub extraction
        github_patterns = [
            r'github\.com/[A-Za-z0-9-]+',
            r'www\.github\.com/[A-Za-z0-9-]+',
            r'https?://(?:www\.)?github\.com/[A-Za-z0-9-]+'
        ]
        
        for pattern in github_patterns:
            github_match = re.search(pattern, text, re.IGNORECASE)
            if github_match:
                contact_info['github'] = github_match.group()
                break
        
        # Website extraction
        website_pattern = r'https?://[A-Za-z0-9.-]+\.[A-Za-z]{2,}'
        website_matches = re.findall(website_pattern, text)
        if website_matches:
            # Filter out LinkedIn and GitHub URLs
            websites = [url for url in website_matches 
                       if 'linkedin.com' not in url and 'github.com' not in url]
            if websites:
                contact_info['website'] = websites[0]
        
        # Location extraction (basic patterns)
        location_patterns = [
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, State
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, Country
        ]
        
        for pattern in location_patterns:
            location_match = re.search(pattern, text)
            if location_match:
                contact_info['location'] = location_match.group(1)
                break
        
        return contact_info
    
    async def _store_file(self, file_content: BinaryIO, filename: str) -> str:
        """Store file in MinIO object storage."""
        try:
            file_content.seek(0)
            file_size = len(file_content.read())
            file_content.seek(0)
            
            # Generate unique filename
            import uuid
            from datetime import datetime
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            unique_id = str(uuid.uuid4())[:8]
            stored_filename = f"{timestamp}_{unique_id}_{filename}"
            
            # Upload file
            self.minio_client.put_object(
                settings.minio_bucket_name,
                stored_filename,
                file_content,
                file_size
            )
            
            return stored_filename
        
        except Exception as e:
            print(f"Warning: Could not store file in MinIO: {e}")
            return filename  # Return original filename if storage fails
    
    async def get_stored_file(self, filename: str) -> Optional[bytes]:
        """Retrieve stored file from MinIO."""
        try:
            response = self.minio_client.get_object(settings.minio_bucket_name, filename)
            return response.read()
        except Exception as e:
            print(f"Could not retrieve file {filename}: {e}")
            return None
    
    async def list_stored_files(self, prefix: str = "") -> List[Dict[str, any]]:
        """List stored files with metadata."""
        try:
            objects = self.minio_client.list_objects(
                settings.minio_bucket_name, 
                prefix=prefix, 
                recursive=True
            )
            
            files = []
            for obj in objects:
                files.append({
                    'filename': obj.object_name,
                    'size': obj.size,
                    'last_modified': obj.last_modified,
                    'etag': obj.etag
                })
            
            return files
        except Exception as e:
            print(f"Could not list files: {e}")
            return []
